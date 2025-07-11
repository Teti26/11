import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import json
import re
import io

def interpret_beck(score):
    if score == "" or score is None:
        return ""
    score = int(score)
    if score <= 9:
        return "удовлетворительное эмоциональное состояние"
    elif score <= 19:
        return "лёгкое депрессивное состояние"
    elif score <= 22:
        return "умеренное депрессивное состояние"
    else:
        return "выраженное депрессивное состояние"

def interpret_hads_anxiety(score):
    if score == "" or score is None:
        return ""
    score = int(score)
    if score <= 7:
        return "нормативный уровень тревоги"
    elif score <= 10:
        return "субклинически выраженная тревога"
    else:
        return "клинически выраженная тревога"

def interpret_hads_depression(score):
    if score == "" or score is None:
        return ""
    score = int(score)
    if score <= 7:
        return "нормативный уровень депрессии"
    elif score <= 10:
        return "субклинически выраженное депрессивное состояние"
    else:
        return "клинически выраженное депрессивное состояние"

def interpret_spielberger_state(score):
    if score == "" or score is None:
        return ""
    score = int(score)
    if score <= 30:
        return "низкий уровень тревожности (норма)"
    elif score <= 45:
        return "средний уровень тревожности"
    else:
        return "высокий уровень тревожности"

def interpret_spielberger_trait(score):
    if score == "" or score is None:
        return ""
    score = int(score)
    if score <= 30:
        return "низкий уровень тревожности (норма)"
    elif score <= 45:
        return "средний уровень тревожности"
    else:
        return "высокий уровень тревожности"

def interpret_prikhojan(score):
    if score == "" or score is None:
        return ""
    score = int(score)
    if score <= 2:
        return "тревожность не выражена или чрезмерное спокойствие защитного характера"
    elif score <= 6:
        return "нормативный уровень тревожности"
    elif score <= 8:
        return "несколько завышенная тревожность"
    elif score == 9:
        return "явно повышенная тревожность"
    else:
        return "очень высокая тревожность (группа риска)"

st.set_page_config(page_title="Генератор заключений", layout="centered")

TEMPLATE_MAP = {
    "🧠 Консультативное заключение": {
        "template": "Шаблон_с_интерпретациями.docx",
        "json": "structured_choice_options_консультация_обновлён.json",
        "sections": "structured_sections_консультация_обновлён.json"
    }
}

st.title("🧾 Генератор заключений")
selected_type = st.radio("### Выберите тип заключения:", list(TEMPLATE_MAP.keys()))

TEMPLATE_PATH = TEMPLATE_MAP[selected_type]["template"]
CHOICES_PATH = TEMPLATE_MAP[selected_type]["json"]
SECTIONS_PATH = TEMPLATE_MAP[selected_type].get("sections")

@st.cache_data
def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

@st.cache_resource
def load_template(path):
    return Document(path)

choices = load_json(CHOICES_PATH)
doc_template = load_template(TEMPLATE_PATH)
sections = load_json(SECTIONS_PATH) if SECTIONS_PATH else None

def extract_ordered_markers(doc):
    markers = []
    seen = set()
    for para in doc.paragraphs:
        found = re.findall(r"\{\{(.+?)\}\}", para.text)
        for marker in found:
            if marker not in seen:
                seen.add(marker)
                markers.append(marker)
    return markers

user_values = {}

def render_field(marker):
    quantitative_markers = {
        "шкала_бека": interpret_beck,
        "hads_тревога": interpret_hads_anxiety,
        "hads_депрессия": interpret_hads_depression,
        "спилбергер_ситуативная": interpret_spielberger_state,
        "спилбергер_личностная": interpret_spielberger_trait,
        "прихожан_общая": interpret_prikhojan,
        "прихожан_школьная": interpret_prikhojan,
        "прихожан_самооценочная": interpret_prikhojan,
        "прихожан_межличностная": interpret_prikhojan,
        "прихожан_магическая": interpret_prikhojan,
    }

    if marker in quantitative_markers:
        label = choices.get(marker, {}).get("label", marker)
        raw = st.text_input(label, key=marker)
        if raw and raw.isdigit() and int(raw) > 0:
            score = int(raw)
            user_values[marker] = score
            interp = quantitative_markers[marker](score)
            user_values[marker + "_интерпретация"] = interp
            st.info(interp)
        else:
            user_values[marker] = ""
            user_values[marker + "_интерпретация"] = ""
        return

    field = choices.get(marker)
    label = field.get("label", marker.replace("_", " ").capitalize()) if field else marker
    help_text = field.get("help", "") if field else ""

    if not field:
        user_values[marker] = st.text_input(label, key=marker)
        return

    if marker in ("отмечается", "отмечается_семья", "класс", "вариант_программы"):
        return
    if marker == "уточнение" and "уточнение" not in user_values:
        user_values["уточнение"] = ""

    field = choices.get(marker)
    label = field.get("label", marker.replace("_", " ").capitalize()) if field else marker
    help_text = field.get("help", "") if field else ""

    if not field:
        user_values[marker] = st.text_input(label, key=marker)
        return

    if field["type"] == "radio":
        selected = st.radio(label, field["options"], help=help_text, key=marker, horizontal=True)
        user_values[marker] = selected

        if marker == "стратегии":
            if selected == "Выявляются признаки дисфункциональных стратегий семейного взаимодействия.":
                subfield = choices.get("отмечается_семья")
                if subfield:
                    sublabel = subfield.get("label", "отмечается")
                    subhelp = subfield.get("help", "")
                    selected_items = st.multiselect(sublabel, subfield["options"], help=subhelp, key="отмечается_семья_dyn")
                    user_values["отмечается_семья"] = ", ".join(selected_items)
            else:
                user_values["отмечается_семья"] = ""

        if marker == "обучается_в":
            if selected in ("школе", "коррекционной школе"):
                subfield = choices.get("класс")
                if subfield:
                    sublabel = subfield.get("label", "Класс")
                    subhelp = subfield.get("help", "")
                    selected_class = st.radio(sublabel, subfield["options"], help=subhelp, key="класс_dyn", horizontal=True)
                    user_values["класс"] = selected_class
            else:
                user_values["класс"] = ""

        if marker == "школьная_программа":
            if selected == "по адаптированной программе":
                variant = st.text_input("вариант программы", key="вариант_программы_dyn")
                user_values["вариант_программы"] = variant
            else:
                user_values["вариант_программы"] = ""

        if marker == "речевой_контакт" and selected == "недоступен в связи с":
            уточнение = st.text_input("Уточнение причины недоступности речевого контакта", key="уточнение")
            user_values["уточнение"] = уточнение

    elif field["type"] == "multiselect":
        if marker == "рекомендации":
            st.markdown("### ✅ Выбор рекомендаций")
            selected_items = []
            for option in field["options"]:
                if st.checkbox(option, key=f"{marker}_{option}"):
                    selected_items.append(option)
            user_values[marker] = "\n".join(selected_items)
        else:
            selected = st.multiselect(label, field["options"], help=help_text, key=marker)
            user_values[marker] = ", ".join(selected)

    elif field["type"] == "text":
        if marker == "по_данным_проективных_методик":
            user_values[marker] = st.text_area(label, help=help_text, key=marker, height=200)
        else:
            user_values[marker] = st.text_input(label, help=help_text, key=marker)

    elif field["type"] == "number":
        min_val = field.get("min", 0)
        max_val = field.get("max", 100)
        score = st.number_input(label, min_val, max_val, help=help_text, key=marker)
        user_values[marker] = score
        func_name = field.get("interpretation_function")
        if func_name:
            func = globals().get(func_name)
            if func:
                interp = func(score)
                user_values[f"{marker}_интерпретация"] = interp
                st.info(interp)

def render_preview(doc, values):
    preview = ""
    quant_keys = [
        "шкала_бека", "шкала_бека_интерпретация",
        "hads_тревога", "hads_тревога_интерпретация",
        "hads_депрессия", "hads_депрессия_интерпретация",
        "спилбергер_ситуативная", "спилбергер_ситуативная_интерпретация",
        "спилбергер_личностная", "спилбергер_личностная_интерпретация",
        "прихожан_общая", "прихожан_общая_интерпретация",
        "прихожан_школьная", "прихожан_школьная_интерпретация",
        "прихожан_самооценочная", "прихожан_самооценочная_интерпретация",
        "прихожан_межличностная", "прихожан_межличностная_интерпретация",
        "прихожан_магическая", "прихожан_магическая_интерпретация"
    ]
    for para in doc.paragraphs:
        text = para.text
        skip_para = False
        for key in re.findall(r"\{\{(.+?)\}\}", text):
            if key in quant_keys and (not values.get(key)):
                skip_para = True
        if skip_para:
            continue
        for key in re.findall(r"\{\{(.+?)\}\}", text):
            val = values.get(key, "")
            text = text.replace(f"{{{{{key}}}}}", str(val or ""))
        preview += text.strip() + "\n\n"
    return preview.strip()

def fill_template(doc, values):
    new_doc = Document()
    quant_keys = [
        "шкала_бека", "шкала_бека_интерпретация",
        "hads_тревога", "hads_тревога_интерпретация",
        "hads_депрессия", "hads_депрессия_интерпретация",
        "спилбергер_ситуативная", "спилбергер_ситуативная_интерпретация",
        "спилбергер_личностная", "спилбергер_личностная_интерпретация",
        "прихожан_общая", "прихожан_общая_интерпретация",
        "прихожан_школьная", "прихожан_школьная_интерпретация",
        "прихожан_самооценочная", "прихожан_самооценочная_интерпретация",
        "прихожан_межличностная", "прихожан_межличностная_интерпретация",
        "прихожан_магическая", "прихожан_магическая_интерпретация"
    ]
    for para in doc.paragraphs:
        text = para.text
        markers = re.findall(r"\{\{(.+?)\}\}", text)
        if any(m in quant_keys for m in markers):
            if all(not str(values.get(m, "")) for m in markers if m in quant_keys):
                continue

        extra_recommendations = []
        for key in markers:
            val = values.get(key, "")
            if key in quant_keys and not val:
                val = ""
            elif key == "класс" and val:
                val = f"{val} классе"
            elif key == "класс":
                val = ""
            elif key.startswith("отмеча") and val:
                val = "Отмечаются: " + val + "."
            elif key.startswith("отмеча"):
                val = ""
            elif key == "уточнение" and val:
                val = f" ({val})"
            elif key == "рекомендации" and val:
                val = val.split("\n")
                extra_recommendations = val
                val = ""
            else:
                val = str(val or "")
            text = text.replace(f"{{{{{key}}}}}", val if isinstance(val, str) else str(val))

        p = new_doc.add_paragraph()
        run = p.add_run(text.strip())
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for item in extra_recommendations:
            if item.strip():
                rec_p = new_doc.add_paragraph()
                rec_run = rec_p.add_run(item.strip())
                rec_run.font.name = 'Times New Roman'
                rec_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                rec_run.font.size = Pt(12)
                rec_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return new_doc

# ====== ОСНОВНОЙ ИНТЕРФЕЙС ======
st.markdown(f"<h2 style='text-align:center'>{selected_type}</h2>", unsafe_allow_html=True)
if sections:
    for section_title, field_list in sections.items():
        with st.expander(section_title, expanded=True):
            for marker in field_list:
                render_field(marker)
else:
    for marker in extract_ordered_markers(doc_template):
        render_field(marker)

if st.button("👁 Предпросмотр"):
    st.text_area("Предпросмотр", render_preview(doc_template, user_values), height=300)

if st.button("📄 Сгенерировать документ"):
    filled = fill_template(doc_template, user_values)
    buffer = io.BytesIO()
    filled.save(buffer)
    buffer.seek(0)
    st.download_button(
        "📥 Скачать .docx",
        data=buffer,
        file_name=f"{selected_type.split()[1].lower()}_заключение.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ====== ЖИРАФ В САМОМ КОНЦЕ СТРАНИЦЫ ======
st.markdown("---")
st.image("giraffe.png", caption="", use_container_width=True)
